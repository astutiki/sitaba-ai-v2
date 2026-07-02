"""
Export API
AI SINTA
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

from datetime import datetime
import os
import json
import csv

from docx import Document
from openpyxl import Workbook
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
)
from reportlab.lib.styles import getSampleStyleSheet


router = APIRouter(
    prefix="/export",
    tags=["Export"]
)

EXPORT_FOLDER = "exports"

os.makedirs(
    EXPORT_FOLDER,
    exist_ok=True
)


# =====================================================
# MODEL
# =====================================================

class ExportRequest(BaseModel):

    question: str

    answer: str

    format: str = "pdf"


# =====================================================
# PDF
# =====================================================

def export_pdf(question, answer):

    filename = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    styles = getSampleStyleSheet()

    doc = SimpleDocTemplate(filepath)

    story = [

        Paragraph("<b>AI SINTA</b>", styles["Heading1"]),

        Paragraph("<b>Pertanyaan</b>", styles["Heading2"]),

        Paragraph(question, styles["BodyText"]),

        Paragraph("<b>Jawaban</b>", styles["Heading2"]),

        Paragraph(answer, styles["BodyText"]),

        Paragraph(
            datetime.now().strftime("%d-%m-%Y %H:%M:%S"),
            styles["Italic"],
        )

    ]

    doc.build(story)

    return filepath


# =====================================================
# WORD
# =====================================================

def export_word(question, answer):

    filename = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    doc = Document()

    doc.add_heading(
        "AI SINTA",
        level=1
    )

    doc.add_heading(
        "Pertanyaan",
        level=2
    )

    doc.add_paragraph(question)

    doc.add_heading(
        "Jawaban",
        level=2
    )

    doc.add_paragraph(answer)

    doc.add_paragraph(
        datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    )

    doc.save(filepath)

    return filepath


# =====================================================
# EXCEL
# =====================================================

def export_excel(question, answer):

    filename = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    wb = Workbook()

    ws = wb.active

    ws.title = "AI SINTA"

    ws.append([
        "Pertanyaan",
        "Jawaban",
        "Tanggal"
    ])

    ws.append([
        question,
        answer,
        datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    ])

    wb.save(filepath)

    return filepath


# =====================================================
# CSV
# =====================================================

def export_csv(question, answer):

    filename = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    with open(
        filepath,
        "w",
        newline="",
        encoding="utf-8"
    ) as f:

        writer = csv.writer(f)

        writer.writerow([
            "Pertanyaan",
            "Jawaban",
            "Tanggal"
        ])

        writer.writerow([
            question,
            answer,
            datetime.now().strftime("%d-%m-%Y %H:%M:%S")
        ])

    return filepath


# =====================================================
# TXT
# =====================================================

def export_txt(question, answer):

    filename = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    with open(
        filepath,
        "w",
        encoding="utf-8"
    ) as f:

        f.write("AI SINTA\n\n")

        f.write("PERTANYAAN\n")

        f.write(question)

        f.write("\n\n")

        f.write("JAWABAN\n")

        f.write(answer)

    return filepath


# =====================================================
# JSON
# =====================================================

def export_json(question, answer):

    filename = f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    data = {

        "question": question,

        "answer": answer,

        "created_at": datetime.now().isoformat()

    }

    with open(
        filepath,
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            data,
            f,
            indent=4,
            ensure_ascii=False
        )

    return filepath


# =====================================================
# EXPORT
# =====================================================

@router.post("/")
def export_chat(data: ExportRequest):

    format = data.format.lower()

    if format == "pdf":

        path = export_pdf(
            data.question,
            data.answer
        )

    elif format == "docx":

        path = export_word(
            data.question,
            data.answer
        )

    elif format == "xlsx":

        path = export_excel(
            data.question,
            data.answer
        )

    elif format == "csv":

        path = export_csv(
            data.question,
            data.answer
        )

    elif format == "txt":

        path = export_txt(
            data.question,
            data.answer
        )

    elif format == "json":

        path = export_json(
            data.question,
            data.answer
        )

    else:

        raise HTTPException(
            status_code=400,
            detail="Format export tidak didukung."
        )

    return {

        "success": True,

        "message": "Export berhasil.",

        "file": os.path.basename(path),

        "download": f"/export/download/{os.path.basename(path)}"

    }


# =====================================================
# DOWNLOAD
# =====================================================

@router.get("/download/{filename}")
def download_file(filename: str):

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    if not os.path.exists(filepath):

        raise HTTPException(
            status_code=404,
            detail="File tidak ditemukan."
        )

    return FileResponse(
        filepath,
        filename=filename
    )


# =====================================================
# LIST FILE
# =====================================================

@router.get("/files")
def list_export():

    files = []

    for file in os.listdir(EXPORT_FOLDER):

        files.append({

            "filename": file,

            "download": f"/export/download/{file}"

        })

    return {

        "success": True,

        "total": len(files),

        "files": files

    }


# =====================================================
# DELETE FILE
# =====================================================

@router.delete("/{filename}")
def delete_export(filename: str):

    filepath = os.path.join(
        EXPORT_FOLDER,
        filename
    )

    if not os.path.exists(filepath):

        raise HTTPException(
            status_code=404,
            detail="File tidak ditemukan."
        )

    os.remove(filepath)

    return {

        "success": True,

        "message": "File berhasil dihapus."

    }